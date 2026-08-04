"""
Extracts raw text from an uploaded contract file.

Strategy per file type:
- pdf: try text-layer extraction (pdfplumber) page by page. If a page yields
  next to no text (i.e. it's a scanned image), fall back to rasterizing that
  page and running Tesseract OCR on it. This handles mixed documents (some
  real pages, some scanned pages) correctly instead of guessing up front.
- docx: python-docx paragraph + table extraction.
- png/jpg/jpeg: straight to Tesseract OCR.

Raises ExtractionError on failure so the caller can mark the contract
status="failed" with a real reason instead of silently storing empty text.
"""
import io
import logging

logger = logging.getLogger("clauseiq.extraction")

MIN_CHARS_PER_PAGE_BEFORE_OCR_FALLBACK = 20


class ExtractionError(Exception):
    pass


def extract_text(file_bytes: bytes, file_type: str) -> str:
    file_type = file_type.lower()
    try:
        if file_type == "pdf":
            return _extract_pdf(file_bytes)
        if file_type in ("docx", "doc"):
            return _extract_docx(file_bytes)
        if file_type in ("png", "jpg", "jpeg"):
            return _extract_image(file_bytes)
    except ExtractionError:
        raise
    except Exception as exc:  # noqa: BLE001 — convert any library-specific error to our own type
        logger.exception("Extraction failed for file_type=%s", file_type)
        raise ExtractionError(f"Could not extract text from {file_type} file: {exc}") from exc

    raise ExtractionError(f"Unsupported file type for extraction: {file_type}")


def _extract_pdf(file_bytes: bytes) -> str:
    import pdfplumber

    pages_text: list[str] = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page_num, page in enumerate(pdf.pages):
            text = (page.extract_text() or "").strip()
            if len(text) < MIN_CHARS_PER_PAGE_BEFORE_OCR_FALLBACK:
                text = _ocr_pdf_page(file_bytes, page_num) or text
            pages_text.append(text)

    full_text = "\n\n".join(pages_text).strip()
    if not full_text:
        raise ExtractionError("No extractable text found in PDF (empty or corrupt file)")
    return full_text


def _ocr_pdf_page(file_bytes: bytes, page_num: int) -> str:
    """Rasterize a single PDF page and OCR it. Used as a per-page fallback
    for scanned/image-only pages inside an otherwise-text PDF."""
    from pdf2image import convert_from_bytes
    import pytesseract

    images = convert_from_bytes(
        file_bytes, first_page=page_num + 1, last_page=page_num + 1, dpi=300
    )
    if not images:
        return ""
    return pytesseract.image_to_string(images[0]).strip()


def _extract_docx(file_bytes: bytes) -> str:
    import docx

    document = docx.Document(io.BytesIO(file_bytes))
    parts: list[str] = [p.text for p in document.paragraphs if p.text.strip()]

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))

    full_text = "\n".join(parts).strip()
    if not full_text:
        raise ExtractionError("No extractable text found in DOCX (empty file)")
    return full_text


def _extract_image(file_bytes: bytes) -> str:
    from PIL import Image
    import pytesseract

    image = Image.open(io.BytesIO(file_bytes))
    text = pytesseract.image_to_string(image).strip()
    if not text:
        raise ExtractionError("No text detected in image via OCR")
    return text
